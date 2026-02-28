"""
AI-Powered Resume Screening & Ranking System
Flask Backend — Main Application
"""

import os
import io
import csv
import json
import ast
import re
import pandas as pd
from flask import Flask, render_template, request, jsonify, send_file, session
from werkzeug.utils import secure_filename

# PDF extraction
from pdfminer.high_level import extract_text as pdf_extract_text
from pdfminer.layout import LAParams

# DOCX extraction
import docx as python_docx

# Scoring
from utils.scoring import rank_candidates_from_csv, rank_candidates_from_uploads, get_vectorizer
from utils.skill_extraction import extract_skills_from_text

# PDF export
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.enums import TA_CENTER, TA_LEFT

app = Flask(__name__)
app.secret_key = os.urandom(24)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload

DATASET_PATH = os.path.join(os.path.dirname(__file__), 'dataset', 'resumes.csv')
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}

# ── Dataset ──────────────────────────────────────────────────────────────────

def load_dataset():
    """Load and clean the CSV dataset."""
    try:
        df = pd.read_csv(DATASET_PATH, encoding='utf-8-sig')
        # Rename BOM column if present
        df.columns = [c.strip().lstrip('\ufeff') for c in df.columns]
        return df
    except Exception as e:
        print(f"Dataset load error: {e}")
        return pd.DataFrame()


# Pre-load dataset at startup
print("Loading dataset...")
DATASET = load_dataset()
print(f"Dataset loaded: {len(DATASET)} records")


# ── File Helpers ──────────────────────────────────────────────────────────────

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(file_bytes):
    """Extract text from PDF bytes using pdfminer."""
    try:
        pdf_file = io.BytesIO(file_bytes)
        laparams = LAParams(line_margin=0.5, word_margin=0.1)
        text = pdf_extract_text(pdf_file, laparams=laparams)
        return text or ""
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return ""


def extract_text_from_docx(file_bytes):
    """Extract text from DOCX bytes."""
    try:
        doc = python_docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return '\n'.join(paragraphs)
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return ""


def extract_text_from_file(file):
    """Route to correct extractor based on file extension."""
    filename = secure_filename(file.filename)
    ext = filename.rsplit('.', 1)[1].lower()
    file_bytes = file.read()

    if ext == 'pdf':
        return extract_text_from_pdf(file_bytes)
    elif ext == 'docx':
        return extract_text_from_docx(file_bytes)
    elif ext == 'txt':
        return file_bytes.decode('utf-8', errors='ignore')
    return ""


# ── Routes ────────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    """Main recruiter interface."""
    return render_template('index.html')


@app.route('/api/evaluate', methods=['POST'])
def evaluate():
    """
    Main evaluation endpoint.
    Handles both CSV dataset screening and file uploads.
    """
    mode = request.form.get('mode', 'csv')  # 'csv' or 'upload'
    job_description = request.form.get('job_description', '').strip()
    sample_size = int(request.form.get('sample_size', 50))

    if not job_description:
        return jsonify({'error': 'Job description is required'}), 400

    results = []

    if mode == 'csv':
        if DATASET.empty:
            return jsonify({'error': 'Dataset not available'}), 500

        # Optional: filter by job role keyword
        role_filter = request.form.get('role_filter', '').strip()
        df = DATASET.copy()
        if role_filter:
            mask = df.apply(
                lambda row: role_filter.lower() in str(row.get('job_position_name', '')).lower() or
                            role_filter.lower() in str(row.get('positions', '')).lower(),
                axis=1
            )
            df = df[mask] if mask.any() else df

        results = rank_candidates_from_csv(df, job_description, sample_size=sample_size)

    elif mode == 'upload':
        files = request.files.getlist('resumes')
        if not files or all(f.filename == '' for f in files):
            return jsonify({'error': 'No files uploaded'}), 400

        files_data = []
        for f in files:
            if f and allowed_file(f.filename):
                text = extract_text_from_file(f)
                if text.strip():
                    name = secure_filename(f.filename).rsplit('.', 1)[0].replace('_', ' ').replace('-', ' ').title()
                    files_data.append({'name': name, 'text': text})

        if not files_data:
            return jsonify({'error': 'Could not extract text from uploaded files'}), 400

        results = rank_candidates_from_uploads(files_data, job_description)

    # Store in session for export
    session['last_results'] = results
    session['last_jd'] = job_description

    return jsonify({
        'results': results,
        'total': len(results),
        'jd_skills': extract_skills_from_text(job_description)[:20],
        'mode': mode,
    })


@app.route('/api/export/csv')
def export_csv():
    """Export results as CSV."""
    results = session.get('last_results', [])
    if not results:
        return jsonify({'error': 'No results to export'}), 400

    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=[
        'rank', 'name', 'score', 'fit_label',
        'matched_skills', 'missing_skills'
    ])
    writer.writeheader()
    for r in results:
        writer.writerow({
            'rank': r.get('rank', ''),
            'name': r.get('name', ''),
            'score': f"{r.get('score', 0):.1f}%",
            'fit_label': r.get('fit_label', ''),
            'matched_skills': ', '.join(r.get('matched_skills', [])),
            'missing_skills': ', '.join(r.get('missing_skills', [])),
        })

    output.seek(0)
    return send_file(
        io.BytesIO(output.getvalue().encode()),
        mimetype='text/csv',
        as_attachment=True,
        download_name='screening_results.csv'
    )


@app.route('/api/export/pdf')
def export_pdf():
    """Export results as PDF report."""
    results = session.get('last_results', [])
    jd = session.get('last_jd', '')
    if not results:
        return jsonify({'error': 'No results to export'}), 400

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        rightMargin=1.5*cm, leftMargin=1.5*cm,
        topMargin=2*cm, bottomMargin=2*cm
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Title'],
                                  fontSize=18, textColor=colors.HexColor('#00d9c4'),
                                  spaceAfter=6)
    subtitle_style = ParagraphStyle('Sub', parent=styles['Normal'],
                                     fontSize=9, textColor=colors.HexColor('#8b949e'),
                                     spaceAfter=12)
    header_style = ParagraphStyle('Header', parent=styles['Heading2'],
                                   fontSize=11, textColor=colors.HexColor('#e6edf3'),
                                   spaceAfter=6)

    story = []
    story.append(Paragraph("Resume Screening Report", title_style))
    story.append(Paragraph(f"Job Description Summary: {jd[:200]}{'...' if len(jd) > 200 else ''}", subtitle_style))
    story.append(Spacer(1, 0.4*cm))

    # Build table data
    headers = ['Rank', 'Candidate', 'Score', 'Fit Category', 'Matched Skills', 'Missing Skills']
    table_data = [headers]

    for r in results[:50]:  # Cap at 50 rows
        table_data.append([
            str(r.get('rank', '')),
            r.get('name', '')[:35],
            f"{r.get('score', 0):.1f}%",
            r.get('fit_label', ''),
            ', '.join(r.get('matched_skills', []))[:80] or '—',
            ', '.join(r.get('missing_skills', []))[:80] or '—',
        ])

    col_widths = [1.2*cm, 4.5*cm, 1.8*cm, 2.8*cm, 4.5*cm, 4.5*cm]
    table = Table(table_data, colWidths=col_widths, repeatRows=1)

    # Style
    fit_colors = {'Strong Fit': '#22c55e', 'Moderate Fit': '#f59e0b', 'Weak Fit': '#ef4444'}
    table_style = [
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#161b22')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#00d9c4')),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 8),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#0d1117'), colors.HexColor('#111318')]),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor('#c9d1d9')),
        ('FONTSIZE', (0, 1), (-1, -1), 7),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('ALIGN', (0, 0), (0, -1), 'CENTER'),
        ('ALIGN', (2, 0), (2, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 0.3, colors.HexColor('#30363d')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('LEFTPADDING', (0, 0), (-1, -1), 5),
    ]

    # Color fit category cells
    for i, r in enumerate(results[:50], start=1):
        fit = r.get('fit_label', '')
        hex_color = fit_colors.get(fit, '#8b949e')
        table_style.append(('TEXTCOLOR', (3, i), (3, i), colors.HexColor(hex_color)))
        table_style.append(('FONTNAME', (3, i), (3, i), 'Helvetica-Bold'))

    table.setStyle(TableStyle(table_style))
    story.append(table)

    doc.build(story)
    buffer.seek(0)

    return send_file(
        buffer,
        mimetype='application/pdf',
        as_attachment=True,
        download_name='screening_results.pdf'
    )


@app.route('/api/dataset/stats')
def dataset_stats():
    """Return dataset statistics for the UI."""
    if DATASET.empty:
        return jsonify({'total': 0})

    col = 'job_position_name'
    if col not in DATASET.columns:
        col = [c for c in DATASET.columns if 'job' in c.lower() or 'position' in c.lower()]
        col = col[0] if col else None

    roles = []
    if col:
        roles = DATASET[col].dropna().unique().tolist()[:30]

    return jsonify({
        'total': len(DATASET),
        'roles': [str(r) for r in roles if str(r) != 'nan'],
    })


if __name__ == '__main__':
    app.run(debug=True, port=5000)
